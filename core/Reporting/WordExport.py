from docx import Document
from docx.shared import Inches
from copy import deepcopy
import os
import sys
from docx.enum.text import WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL # pylint: disable=no-name-in-module
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH # pylint: disable=no-name-in-module
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_LINE_SPACING # pylint: disable=no-name-in-module
from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml
import docx
from docx.shared import RGBColor
from docx.oxml.text.paragraph import CT_P
from docx.shared import Pt, Cm
from docx.text.paragraph import Paragraph
from docx.oxml.xmlchemy import OxmlElement
import json
from bson.objectid import ObjectId
from datetime import datetime
import requests
import re
import io
from core.Models.Ip import Ip
from core.Models.Port import Port
from core.Models.Defect import Defect
from core.Application.Dialogs.ChildDialogQuestion import ChildDialogQuestion
import tkinter as tk

class DummyProgressBar:
    def update(self):
        return

def downloadImgData(url):
    data = requests.get(url)
    if data.status_code != 200:
        return None
    data = data.content
    data = io.BytesIO(data)
    return data

def getParagraphs(document):
    """ Retourne un generateur pour tous les paragraphes du document.
        La page d'entête n'étant pas incluse dans documents.paragraphs."""
    body = document._body._body # pylint: disable=protected-access
    ps = body.xpath('//w:p')
    for p in ps:
        yield Paragraph(p, document._body) # pylint: disable=protected-access

def set_cell_border(cell, **kwargs):
    """
    Set cell`s border
    Usage:

    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc # pylint: disable=protected-access
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def add_hyperlink(paragraph, text, url, style):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run ()
    r._r.append (hyperlink) # pylint: disable=protected-access

    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it

    r.font.underline = True
    if style is not None:
        paragraph.runs[-1].style = style
    return hyperlink

def set_hyperlink(paragraph, run, url, text, style):
    # This gets access to the document.xml.rels file and gets a new relation id value
    run.font.underline = True
    if style is not None:
        run.style = style
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    index_in_paragraph = paragraph._p.index(run.element)
    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )
    hyperlink.append(run.element)
    paragraph._p[index_in_paragraph:index_in_paragraph] = [hyperlink]
    # Delete this if using a template that has the hyperlink style in it
    return hyperlink

def insert_paragraph_after(paragraph, text=None, style=None):
    """
    Insert a new paragraph after the given paragraph.
        Args:
            paragraph: the paragraph object after which the new paragraph will be created
        
        Optional Args:
            text: a string of text to write in the new paragraph, Default = None
            style: a style to be applied on the new paragraph, Default = None

        Returns:
            Returns a paragraph object for the new added paragraph
    """
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)  # pylint: disable=protected-access
    new_para = Paragraph(new_p, paragraph._parent)  # pylint: disable=protected-access
    if text is not None:
        new_para.add_run(text)
    if style is not None:
        new_para.style = style
    return new_para


def _replaceTextInTable(document, search, replace):
    """
    Replace at a run level a text that will be searched in all table cells.
        Args:
            document: the document object (opened word docx with tables inside)
            search: The string to search for that will be replaced. //!\\ The search is done at a run level
            replace: The string to replace the search by.
    """
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if search in paragraph.text:
                        inline = paragraph.runs
                        # Loop added to work with runs (strings with same style)
                        for i in range(len(inline)):
                            if search in inline[i].text:
                                inline[i].text = inline[i].text.replace(search, replace)

def replaceTextInTable(table, search, replace):
    """
    Replace at a run level a text that will be searched in the given table cells.
        Args:
            table: the table we want to replace a text inside
            search: The string to search for that will be replaced. //!\\ The search is done at a run level
            replace: The string to replace the search by.
        Returns:
            Return None if nothing was found, (row's index, cell's index) otherwise
    """
    for i_row, row in enumerate(table.rows):
        for i_cell, cell in enumerate(row.cells):
            for paragraph in cell.paragraphs:
                if search in paragraph.text:
                    inline = paragraph.runs
                    # Loop added to work with runs (strings with same style)
                    for i in range(len(inline)):
                        if search in inline[i].text:
                            inline[i].text = inline[i].text.replace(search, replace)
                            return i_row, i_cell

def replaceTextInParagraph(paragraph, search, replace):
    """
    Replace at a run level a text that will be searched in the given paragraph.
        Args:
            table: the table we want to replace a text inside
            search: The string to search for that will be replaced. //!\\ The search is done at a run level
            replace: The string to replace the search by.
        Returns:
            Return None if nothing was found, the run where the string was replaced otherwise
    """
    if search in paragraph.text:
        for run in paragraph.runs:
            inline = paragraph.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if search in inline[i].text:
                    inline[i].text = inline[i].text.replace(search, replace)  
                    return run
        paragraph.text = paragraph.text.replace(search, replace)
        return paragraph
    return None

def _replaceTextInParagraphes(document, search, replace):
    """
    Replace at a run level a text that will be searched in all the document's paragraphs.
        Args:
            document: the document object (opened word docx with at lease one paragraph inside)
            search: The string to search for that will be replaced. //!\\ The search is done at a run level
            replace: The string to replace the search by.
    """
    ps = getParagraphs(document)
    for paragraph in ps:
        if search in paragraph.text:
            inline = paragraph.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if search in inline[i].text:
                    inline[i].text = inline[i].text.replace(search, replace)

def _replaceTextInHeaders(document, search, replace):
    """
    Replace at a run level a text that will be searched in all the document's headers.
        Args:
            document: the document object (opened word docx with at lease one paragraph inside)
            search: The string to search for that will be replaced. //!\\ The search is done at a run level
            replace: The string to replace the search by.
    """
    for section in document.sections:
        header = section.header
        for paragraph in header.paragraphs:
            replaceTextInParagraph(paragraph, search, replace)


def replaceTextInDocument(document, search, replace):
    """
    Replace at a run level a text that will be searched in all the document paragraphs and tables.
        Args:
            document: the document object (opened word docx with at lease one paragraph or table inside)
            search: The string to search for that will be replaced. //!\\ The search is done at a run level
            replace: The string to replace the search by.
    """
    _replaceTextInTable(document, search, replace)
    _replaceTextInParagraphes(document, search, replace)
    _replaceTextInHeaders(document, search, replace)

def findRowContaining(document, search):
    """
    Find the first row of a table in all the given document to contain the given search text.
        Args:
            document: the document object (opened word docx with at least table inside)
            search: The string to search for that will be replaced. //!\\ The search is done at a paragraph level
        Returns:
            Return (None, None) if the text was not found, else, it returns the tuple (table object, table index in document)
    """
    for table_i, table in enumerate(document.tables):
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if search in paragraph.text:
                        return table, table_i
    return None, None

def findParagraphContaining(document, search):
    """
    Find the first paragraph in all the given document to contain the given search text.
        Args:
            document: the document object (opened word docx with at least one paragraph inside)
            search: The string to search for that will be replaced. //!\\ The search is done at a paragraph level
        Returns:
            Return None if the text was not found, else, it returns the paragraph object where the text was found
    """
    ps = getParagraphs(document)
    for paragraph in ps:
        if search in paragraph.text:
            return paragraph
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if search in paragraph.text:
                        return paragraph
    return None


def delete_paragraph(paragraph):
    """
    Delete a paragraph.
        Args:
            paragraph: the paragraph object to delete.
    """
    p = paragraph._element  # pylint: disable=protected-access
    try:
        p.getparent().remove(p)
    except Exception: # pylint: disable=broad-except
        print("No parent found for element "+str(p))
    p._p = p._element = None  # pylint: disable=protected-access

def write_res_table_defect_line(table_res, risks_font_colors, risks_bg_colors, count, o_defect):
    new_row_cells = table_res.add_row().cells
    level = o_defect["risk"]
    fill_cell(new_row_cells[0], "D"+str(count), risks_font_colors[level], risks_bg_colors[level], True)
    run = new_row_cells[0].paragraphs[0].runs[0]
    font = run.font
    run.bold = True
    font.size = Pt(11)
    fill_cell(new_row_cells[1], o_defect["title"], risks_font_colors[level], risks_bg_colors[level])
    run = new_row_cells[1].paragraphs[0].runs[0]
    font = run.font
    font.size = Pt(11)
    fill_cell(new_row_cells[3], "Risque\n", risks_font_colors[level], risks_bg_colors[level])
    new_row_cells[3].paragraphs[0].add_run(o_defect["risk"], style="contentStyle")
    new_row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = new_row_cells[3].paragraphs[0].runs[0]
    font = header_run.font
    font.size = Pt(9)
    header_run.italic = True
    value_run = new_row_cells[3].paragraphs[0].runs[1]
    font = value_run.font
    font.size = Pt(11)
    font.color.rgb = risks_font_colors[o_defect["risk"]]
    fill_cell(new_row_cells[4], "Exploitation\n", risks_font_colors[level], risks_bg_colors[level])
    new_row_cells[4].paragraphs[0].add_run(o_defect["risk"], style="contentStyle")
    new_row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_run = new_row_cells[4].paragraphs[0].runs[0]
    font = header_run.font
    font.size = Pt(9)
    header_run.italic = True
    value_run = new_row_cells[4].paragraphs[0].runs[1]
    font = value_run.font
    font.size = Pt(11)
    font.color.rgb = risks_font_colors[o_defect["risk"]]
    fill_cell(new_row_cells[5], "Impact\n", risks_font_colors[level], risks_bg_colors[level])
    new_row_cells[5].paragraphs[0].add_run(o_defect["impact"], style="contentStyle")
    new_row_cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_run = new_row_cells[5].paragraphs[0].runs[0]
    font = header_run.font
    font.size = Pt(9)
    header_run.italic = True
    value_run = new_row_cells[5].paragraphs[0].runs[1]
    font = value_run.font
    font.size = Pt(11)
    font.color.rgb = risks_font_colors[o_defect["risk"]]

def write_res_table_fix_line(table_res, risk_bg_color, id_correctif, fixe):
    new_row_cells = table_res.add_row().cells
    black = RGBColor(0x26, 0x32, 0x32)
    gray = "E9E9E9"
    lightgray = RGBColor(0x79, 0x79, 0x79)
    fill_cell(new_row_cells[0], " ", black, risk_bg_color)
    fill_cell(new_row_cells[1], id_correctif, black, gray, True)
    run = new_row_cells[1].paragraphs[0].runs[0]
    font = run.font
    run.bold = True
    font.size = Pt(11)
    fill_cell(new_row_cells[2], fixe["title"], black, gray)
    fill_cell(new_row_cells[4], "Mise en œuvre\n", black, gray)
    new_row_cells[4].paragraphs[0].add_run(fixe["execution"], style='contentStyleFix')
    new_row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_run = new_row_cells[4].paragraphs[0].runs[0]
    font = header_run.font
    font.size = Pt(9)
    font.color.rgb = lightgray
    header_run.italic = True
    fill_cell(new_row_cells[5], "Gain\n", black, gray)
    new_row_cells[5].paragraphs[0].add_run(fixe["gain"], style='contentStyleFix')
    new_row_cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_run = new_row_cells[5].paragraphs[0].runs[0]
    font = header_run.font
    font.size = Pt(9)
    font.color.rgb = lightgray
    header_run.italic = True

def format_block_res_table(table_res, start_defect_line_on_res_table, nb_line_res_table):
    cell_title = table_res.cell(start_defect_line_on_res_table, 1)
    cell_title.merge(table_res.cell(start_defect_line_on_res_table, 2)) # Merge defect title
    cell_id = table_res.cell(start_defect_line_on_res_table, 0)
    try:
        cell_id_fix = table_res.cell(start_defect_line_on_res_table+1, 0)
        if start_defect_line_on_res_table+1 != nb_line_res_table-1:
            cell_id_fix.merge(table_res.cell(nb_line_res_table-1, 0)) # Merge defect id
        set_cell_border(cell_id_fix, start={"sz": 0.5, "val": "single", "color": "#000000", "space": "0"},
            end={"sz": 0.5, "val": "single", "color": "#000000", "space": "0"},
            top={"sz": 0, "val": "none", "color": "#000000", "space": "0"},
            bottom={"sz": 0.5, "val": "single", "color": "#000000", "space": "0"},
        )
    except IndexError:
        pass
    for fixe_line in range(start_defect_line_on_res_table+1, nb_line_res_table):
        cell = table_res.cell(fixe_line, 2)
        cell.merge(table_res.cell(fixe_line, 3)) # Merge defect id
    set_cell_border(cell_title, start={"sz": 0, "val": "none", "color": "#000000", "space": "0"},
        end={"sz": 0, "val": "none", "color": "#000000", "space": "0"},
        top={"sz": 0.5, "val": "single", "color": "#000000", "space": "0"},
        bottom={"sz": 0.5, "val": "single", "color": "#000000", "space": "0"},
    )
    cell_risk = table_res.cell(start_defect_line_on_res_table, 3)
    set_cell_border(cell_risk, start={"sz": 0, "val": "none", "color": "#000000", "space": "0"},
        end={"sz": 0.5, "val": "single", "color": "#000000", "space": "0"},
        top={"sz": 0.5, "val": "single", "color": "#000000", "space": "0"},
        bottom={"sz": 0.5, "val": "single", "color": "#000000", "space": "0"},
    )
    set_cell_border(cell_id, start={"sz": 0.5, "val": "single", "color": "#000000", "space": "0"},
        end={"sz": 0, "val": "none", "color": "#000000", "space": "0"},
        top={"sz": 0.5, "val": "single", "color": "#000000", "space": "0"},
        bottom={"sz": 0, "val": "none", "color": "#000000", "space": "0"},
    )

def populate_defect_summary_table(document, defects_dict):
    """
    Fill a table found with a decorator var_dsum_colId in a line.
    This table will contains a summary of the defects given the excel.
    The template must have a table with var_dsum_colId, var_dsum_colTit, var_dsum_colEase, var_dsum_colImpact, var_dsum_colType
    The template must have another table next to the previous with var_csum_colId, var_csum_colTit, var_csum_colEase, var_csum_colGain
        Args:
            document: the document object (opened word docx with at least one paragraph inside)
            defects_dict: the dictionary of defect gotten in the getDefectDictFromExcel function
    """
    table, table_i = findRowContaining(document, "var_dsum_colId")
    table_c, table_c_i = findRowContaining(document, "var_csum_colId")
    table_res, table_res_i = findRowContaining(document, "var_res_tab")
    if table_i is None:
        raise KeyError("No table row containing var_dsum_colId found")
    if table_c_i is None:
        raise KeyError("No table row containing var_csum_colId found")
    if table_res_i is not None:
        remove_row(table_res, table_res.rows[0])

    _replaceTextInTable, cell_id = replaceTextInTable(table, "var_dsum_colId", "ID")
    _replaceTextInTable, cell_tit = replaceTextInTable(table, "var_dsum_colTit", "Libellé")
    _, cell_ease = replaceTextInTable(table, "var_dsum_colEase", "Exploitation")
    _, cell_impact = replaceTextInTable(table, "var_dsum_colImpact", "Impact")
    _replaceTextInParagraphes, cell_type = replaceTextInTable(table, "var_dsum_colType", "Type")

    _, cell_c_id = replaceTextInTable(table_c, "var_csum_colId", "ID")
    _, cell_c_tit = replaceTextInTable(table_c, "var_csum_colTit", "Libellé")
    _, cell_c_ease = replaceTextInTable(table_c, "var_csum_colEase", "Mise en\nœuvre")
    _, cell_c_gain = replaceTextInTable(table_c, "var_csum_colGain", "Gain en sécurité")
    strong_color = '002060'
    medium_color = '0070C0'
    quick_win_color = '00B0F0'
    white_rgb = RGBColor(0xff, 0xfb, 0xfa)
    black_rgb = RGBColor(0x26, 0x32, 0x32)
    risks_font_colors = {"Critique":white_rgb, "Majeur":white_rgb, "Important":white_rgb, "Mineur":black_rgb}
    risks_bg_colors = {"Critique":"263232", "Majeur":"F8453C", "Important":"EE8200", "Mineur":"FFFF00"}
    fixes_bg_colors = {"Quick Win":quick_win_color, "Faible":quick_win_color, "Modérée":medium_color, "Moyen":medium_color, "Fort":strong_color, "Envergure":strong_color}
    levels = ["Critique", "Majeur", "Important", "Mineur"]
    count = 0
    start_defect_line_on_res_table = 0
    nb_line_res_table = 0
    obj_styles = document.styles
    obj_charstyle = obj_styles.add_style('contentStyle', WD_STYLE_TYPE.CHARACTER) # pylint: disable=no-member
    obj_font = obj_charstyle.font
    obj_font.size = Pt(11)
    obj_font.name = 'Calibri'
    obj_styles = document.styles
    obj_charstyle = obj_styles.add_style('contentStyleFix', WD_STYLE_TYPE.CHARACTER) # pylint: disable=no-member
    obj_font = obj_charstyle.font
    obj_font.size = Pt(11)
    obj_font.name = 'Calibri'
    for level in levels:
        for defect_dict in defects_dict[level].values():
            o_defect = defect_dict["description"]
            new_row_cells = table.add_row().cells
            count += 1
            if table_res_i is not None:
                write_res_table_defect_line(table_res, risks_font_colors, risks_bg_colors, count, o_defect)
            start_defect_line_on_res_table = nb_line_res_table
            nb_line_res_table += 1
            d_types = o_defect["type"]
            defect_types = []
            for types in d_types:
                words_in_type = types.strip().split(" ")
                type_acro = ""
                for word in words_in_type:
                    if word != "":
                        type_acro += word[0].upper()
                defect_types.append(type_acro)
            type_first_letters = ", ".join(defect_types)
            if type_first_letters == "":
                type_first_letters = "N/A"
            fill_cell(new_row_cells[cell_id], "D"+str(count), risks_font_colors[level], risks_bg_colors[level], True)
            fill_cell(new_row_cells[cell_tit], o_defect["title"], risks_font_colors[level], risks_bg_colors[level])
            fill_cell(new_row_cells[cell_ease], o_defect["ease"], risks_font_colors[level], risks_bg_colors[level])
            fill_cell(new_row_cells[cell_impact], o_defect["impact"], risks_font_colors[level], risks_bg_colors[level])
            fill_cell(new_row_cells[cell_type], type_first_letters, risks_font_colors[level], risks_bg_colors[level])
            count_correctif = 1
            id_correctif = "A"+str(count)
            new_row_c_cells = table_c.add_row().cells
            if table_res_i is not None:
                write_res_table_fix_line(table_res, risks_bg_colors[level], id_correctif, {"title":"A definir", "execution":"A def", "gain":"A def"})
                nb_line_res_table += 1
                format_block_res_table(table_res, start_defect_line_on_res_table, nb_line_res_table)
            fill_cell(new_row_c_cells[cell_c_id], id_correctif, None, None, True)
            fill_cell(new_row_c_cells[cell_c_tit], "A definir")
            fill_cell(new_row_c_cells[cell_c_ease], "A def", risks_font_colors["Critique"], fixes_bg_colors["Quick Win"])
            fill_cell(new_row_c_cells[cell_c_gain], "A def", risks_font_colors["Critique"], fixes_bg_colors["Quick Win"])
            if table_res_i is not None:
                format_block_res_table(table_res, start_defect_line_on_res_table, nb_line_res_table)
    # Ajustement de la taille des lignes pour combler la page
    # pageSize = 19 # Valeur totalement arbitraire calculé en regardant si le résultat était bon *PAN*
    # for row in table_c.rows:
    #     row.height = Cm(pageSize/len(table_c.rows))
    # for row in table.rows:
    #     row.height = Cm(pageSize/len(table.rows))
    # Fixe a 1 cm
    for row in table_c.rows:
        row.height = Cm(1)
        row.line_spacing_rule = WD_LINE_SPACING.SINGLE
    for row in table.rows:
        row.height = Cm(1)
        row.line_spacing_rule = WD_LINE_SPACING.SINGLE
def move_table_after(table, paragraph):
    """
    Move a given table after a given paragraph
        Args:
            table: the table to move
            paragraph: the paragraph to put the table after.
    """
    tbl, p = table._tbl, paragraph._p  # pylint: disable=protected-access
    p.addnext(tbl)

def copy_table_after(table, paragraph, addParagraph=None, addPagebreak=False):
    """
    Copy a table and insert it after a given paragraph
        Args:
            table: the table to move
            paragraph: the paragraph to put the table after.
        Optional Args:
            addParagraph: Add another paragraph after the new copied table. default is None.
    """
    tbl, p = table._tbl, paragraph._p  # pylint: disable=protected-access
    new_tbl = deepcopy(tbl)
    new_p = None
    if addParagraph is not None:
        new_p = insert_paragraph_after(paragraph, addParagraph)
        if addPagebreak:
            insertPageBreak(new_p)
    else:
        if addPagebreak:
            insertPageBreak(paragraph)
    p.addnext(new_tbl)
    return new_tbl, new_p

def insertPageBreak(paragraph):
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)

def write_every_defect_fix(fixes, document, last_defect_paragraph, count):
    # print("write fixes for "+str(count)+ " after "+str(last_defect_paragraph.text))
    for fixe_i in range(len(fixes)-1, -1, -1):
        id_correctif = "A"+str(count)
        fixe = fixes[fixe_i]
        if len(fixes) > 1:
            id_correctif += "."+str(fixe_i+1)
        # Find the original fix to copy and fill
        table_c, _ = findRowContaining(document, "var_c_id")
        if table_c is None:
            raise Exception("Error, fixe table was deleted too soon. RIP")
        # Add notes paragraph after it
        _, added_paragraph = copy_table_after(table_c, last_defect_paragraph, "var_c_notes\n")

        copied_table_c, _ = findRowContaining(document, "var_c_id")
        replaceTextInTable(copied_table_c, "var_c_id", id_correctif)
        replaceTextInTable(copied_table_c, "var_c_title", fixe["title"])
        replaceTextInTable(copied_table_c, "var_c_ease", fixe["execution"])
        replaceTextInTable(copied_table_c, "var_c_gain", fixe["gain"])
        desc = fixe["description"].replace("\r", "")
        desc_paras = desc.split("\n")
        fixe_synthesis = fixe.get("synthesis", None)
        if fixe_synthesis is not None:
            desc_paras.insert(0, fixe_synthesis)
        replaceTextInTable(copied_table_c, "var_c_description", desc_paras[0].strip())
        # added_paragraph: premier paragraphe dessous le tableau
        # c_separator: paragraphes supplémentaires éventuelles
        c_separator = added_paragraph
        if len(desc_paras) > 1:
            paragraphAccumulator = ""
            for note in desc_paras[1:]:
                if note.strip() != "":
                    paragraphAccumulator += note.strip()+"\n"
                else:
                    if paragraphAccumulator.strip() != "":
                        c_separator = insert_paragraph_after(c_separator, paragraphAccumulator.strip(), normal_style)
                    paragraphAccumulator = ""
            if paragraphAccumulator.strip() != "":
                c_separator = insert_paragraph_after(c_separator, paragraphAccumulator.strip(), normal_style)
        if added_paragraph == c_separator:
            c_separator.text="\n"
        else:
            delete_paragraph(added_paragraph)
        if fixe_i == len(fixes)-1:
            c_separator = insert_paragraph_after(c_separator, "\n", normal_style)
            insertPageBreak(c_separator)

def insert_images_after(paragraph, pics):
    new_p = insert_paragraph_after(paragraph, "", normal_style)
    new_p.alignement = WD_ALIGN_PARAGRAPH.CENTER
    for pic in pics:
        r = new_p.add_run()
        r.add_picture(pic, width=Cm(17.19))
    return new_p

def write_defect_from_input(result, document, table_d, separator, o_defect, count):
    desc = result[0]["description"].replace("\r", "")
    desc_paras = desc.split("\n")
    synthesis = result[0].get("synthesis", None)
    if synthesis is not None:
        desc_paras.insert(0, synthesis.replace("\r", "").strip())
    desc_paras.insert(1, "ToDo "+str(o_defect["redactor"]))
    desc_paras.insert(2, "\n")
    replaceTextInTable(table_d, "var_d_description", desc_paras[0].strip())
    first_separator = separator
    if len(desc_paras) > 1:
        paragraphAccumulator = ""
        for note in desc_paras[1:]:
            if note.strip() != "":
                paragraphAccumulator += note.strip()+"\n"
            else:
                if paragraphAccumulator.strip() != "":
                    separator = insert_paragraph_after(separator, paragraphAccumulator.strip(), normal_style)
                paragraphAccumulator = ""
        if paragraphAccumulator.strip() != "":
            separator = insert_paragraph_after(separator, paragraphAccumulator.strip(), normal_style)
    details = o_defect.get("details", {})
    for detail in details.values():
        separator = insert_paragraph_after(separator, detail["notes"].strip(), normal_style)
        if detail["pics"]:
            separator = insert_images_after(separator, detail["pics"])
    if first_separator == separator:
        first_separator.text = "\n"
    else:
        delete_paragraph(first_separator)
    write_every_defect_fix(result[0]["fixes"], document, separator, count)



def write_each_defect(document, defects_dict):
    """
    for each default
       Copy a table and a paragraph form the template marked with var_d_id and var_d_separator.
       replace the markers var_d_id var_d_separator var_d_id var_d_title var_d_ease var_d_impact var_d_description
    Then for each fixe of this default
        Copy a table form the template marked with var_c_id.
        replace the markers var_c_id var_d_separator var_c_title var_c_ease var_c_gain var_c_description

        Args:
            document: the document to search elements in
            defects_dict: the dictionary of defect gotten with the dedicated function getDefectDictFromExcel

    """
    levels = ["Critique", "Majeur", "Important", "Mineur"]
    count = 0
    total_len = 0
    for level in levels:
        total_len += len(defects_dict[level].values())
    for level in levels:
        for defect_dict in defects_dict[level].values():
            count += 1
            o_defect = defect_dict["description"]
            table_d, table_i = findRowContaining(document, "var_d_id")
            separator = findParagraphContaining(document, 'var_d_separator')
            copy_table_after(table_d, separator, "var_d_separator")
            replaceTextInTable(table_d, "var_d_id", "D"+str(count))
            replaceTextInTable(table_d, "var_d_title", o_defect["title"])
            replaceTextInTable(table_d, "var_d_ease", o_defect["ease"])
            replaceTextInTable(table_d, "var_d_impact", o_defect["impact"])
            result = [
                    {
                        "id":"0",
                        "title": o_defect["title"],
                        "ease": o_defect["ease"],
                        "impact": o_defect["impact"],
                        "risk": o_defect["risk"],
                        "type": o_defect["type"],
                        "description": "Description",
                        "details": {},
                        "notes": o_defect.get("notes", ""),
                        "fixes": [
                            {
                                "title": "A def",
                                "execution": "Modérée",
                                "gain": "Moyen",
                                "description": "A definir"
                            }
                        ]
                    }
                ]
            o_defect["details"] = {}
            for ids in defect_dict.get("defects_ids", []):
                o_defect["details"][ids] = o_defect["details"].get(ids, {})
                defect_m = Defect.fetchObject({"_id":ObjectId(ids)})
                target = ""
                notes = ""
                if defect_m.ip != "" and defect_m.ip is not None:
                    target += defect_m.ip
                if defect_m.port != "" and defect_m.port is not None:
                    protostr = "" if defect_m.proto == "tcp" else "udp/"
                    target += ":"+protostr+defect_m.port
                if defect_m.notes.strip() == "":
                    defect_m.notes = ""
                elif len(defect_m.notes.strip()) > 2500:
                    defect_m.notes = "ToDo : Too much notes to paste it there."
                if target != "":
                    if defect_m.notes.strip() == "":
                        notes = target+": Vulnerable."
                    else:
                        notes = target+":\n"+defect_m.notes
                else:
                    if defect_m.notes.strip() != "":
                        notes = "Notes: "+defect_m.notes
                o_defect["details"][ids]["notes"] = notes
                pics = []
                for i_proof, _proof in enumerate(defect_m.proofs):
                    ret = defect_m.getProof(i_proof)
                    if ret is not None:
                        pics.append(ret)
                o_defect["details"][ids]["pics"] = pics
            write_defect_from_input(result, document, table_d, separator, o_defect, count)
            progressbar.update()
        # Delete remaining copies
        table_d, table_i = findRowContaining(document, "var_d_id")
        separator = findParagraphContaining(document, 'var_d_separator')
        if separator is not None:
            delete_paragraph(separator)
        for _ in range(len(table_d.rows)):
            remove_row(table_d, table_d.rows[0])
        del document.tables[table_i]
    # Delete remaining correctif copy
    table_c, table_c_i = findRowContaining(document, "var_c_id")
    for _ in range(len(table_c.rows)):
        remove_row(table_c, table_c.rows[0])
    del document.tables[table_c_i]

def remove_row(table, row):
    """
    remove a given row inside a given table
        Args:
            table: the table we want to delete a row in
            row: the row to delete

    """
    tbl = table._tbl  # pylint: disable=protected-access
    tr = row._tr  # pylint: disable=protected-access
    tbl.remove(tr)

def fill_cell(cell, text, font_color=None, bg_color=None, bold=False):
    """
    Fill a table's cell's background with a background color, a text and a font_color for this text
    Also sets the vertical alignement of every cell as centered.
        Args:
            cell: the cell we want to fill
            text: the text to be written inside the cell
        Optional Args:
            font_color: a new font color to use for this text at RGB format (docx rgb). Default is None
            bg_color: A backgroud color to use for the cell at hexa rgb format (FFFFFF is white) default is None.
                        The color is written in xml directly as python-docx does not give a function to do that.
    """
    while len(cell.paragraphs) > 0:
        delete_paragraph(cell.paragraphs[0])
    p = cell.add_paragraph(text)
    p.style = cell_style
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if p.runs:
        p.runs[0].bold = bold
        if font_color is not None:
            p.runs[0].font.color.rgb = font_color
    if bg_color is not None:
        shading_elm_1 = parse_xml((r'<w:shd {} w:fill="'+bg_color+r'"/>').format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading_elm_1)  # pylint: disable=protected-access


def populate_services_table(document, parent):
    table, _ = findRowContaining(document, "var_ssum_ip")
    if table is None:
        return
    replaceTextInTable(table, "var_ssum_ip", "IP")
    replaceTextInTable(table, "var_ssum_port", "Port")
    replaceTextInTable(table, "var_ssum_proto", "Protocole")
    replaceTextInTable(table, "var_ssum_comment", "Commentaires")
    ips = Ip.fetchObjects({"in_scopes": {"$ne":[]}})
    cursorAsList = [c for c in ips] # NOT OPTIMIZED
    nbOfIp = len(cursorAsList)
    if nbOfIp >= 10:
        if parent is None:
            parent = tk.Toplevel()
        dialog = ChildDialogQuestion(parent,
            "Warning", "The service table will contain "+str(nbOfIp)+" IPs with their ports. It would add a lot of pages to the document. Do you want to keep or discard the service table for this report?", ["Keep", "Discard"])
        parent.wait_window(dialog.app)
        if dialog.rvalue == "Discard":
            table._element.getparent().remove(table._element) # pylint: disable=protected-access
            return
    for ip in cursorAsList:
        ligne_deb = len(table.rows)
        ports = Port.fetchObjects({"ip":ip.ip})
        for port in ports:
            new_row_cells = table.add_row().cells
            fill_cell(new_row_cells[1], port.port if port.proto == "tcp" else port.proto+"/"+port.port)
            fill_cell(new_row_cells[2], port.service)
            fill_cell(new_row_cells[3], port.product)

        ligne_end = len(table.rows)
        if len(table.rows) != ligne_deb:
            cell_merge = table.cell(ligne_deb, 0)
            merged_cell = cell_merge.merge(table.cell(ligne_end-1, 0))
            merged_cell.text = ip.ip
            merged_cell.paragraphs[0].style = cell_style
            merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # table.alignment = WD_TABLE_ALIGNMENT.CENTER

def split_run_in_two(paragraph, run, split_index):
    index_in_paragraph = paragraph._p.index(run.element) # pylint: disable=protected-access

    text_before_split = run.text[0:split_index]
    text_after_split = run.text[split_index:]
    
    run.text = text_before_split
    new_run = paragraph.add_run(text_after_split)
    copy_format_manual(run, new_run)
    paragraph._p[index_in_paragraph+1:index_in_paragraph+1] = [new_run.element] # pylint: disable=protected-access
    return [run, new_run]

def split_run_in_three(paragraph, run, split_start, split_end):
    first_split = split_run_in_two(paragraph, run, split_end)
    second_split = split_run_in_two(paragraph, run, split_start)
    return second_split + [first_split[-1]]

def copy_format_manual(runA, runB):
    fontB = runB.font
    fontA = runA.font
    fontB.bold = fontA.bold
    fontB.italic = fontA.italic
    fontB.underline = fontA.underline
    fontB.strike = fontA.strike
    fontB.subscript = fontA.subscript
    fontB.superscript = fontA.superscript
    fontB.size = fontA.size
    fontB.highlight_color = fontA.highlight_color
    fontB.color.rgb = fontA.color.rgb

def markdownHeaderToWordStyle(paragraph, run, style):
    regexHeaders = re.compile(r"^#{1,6} (.+)$", re.MULTILINE)
    matched = re.search(regexHeaders, run.text)
    if matched is not None:
        run.text = matched.group(1)
        paragraph.style = style
    return []

def markdownEmphasisToItalic(paragraph, initialRun):
    """
     de _Man-in-the-Middle_ entre le serveur
     # Absence de support de TLS_FALLBACK_SCSV

    TLS_FALLBACK_SCSV est une option permettant de mitiger les attaques dites de downgrade (type POODLE), afin d’empêcher un attaquant de forcer l’utilisation d’un protocole vulnérable lorsque des protocoles plus récents et sécurisés sont disponibles.

    """
    runs = [initialRun]
    i = 0
    while i < len(runs):
        splitted_runs = splitRunOnMarker(paragraph, runs[i], r"(?<!\w)\*([^\*\n]+)\*(?!\w)", "*")
        if len(splitted_runs) == 3:
            splitted_runs[1].italic = True
            runs.append(splitted_runs[1])
            runs.append(splitted_runs[2])
        else:
            splitted_runs = splitRunOnMarker(paragraph, runs[i], r"(?<!\w)\_([^\_\n]+)\_(?!\w)", "_")
            if len(splitted_runs) == 3:
                splitted_runs[1].italic = True
                runs.append(splitted_runs[1])
                runs.append(splitted_runs[2])
        i+=1
    return runs

def markdownStrongEmphasisToBold(paragraph, initialRun):
    runs = [initialRun]
    i = 0
    while i < len(runs):
        splitted_runs = splitRunOnMarker(paragraph, runs[i], r"(?!\w)\*\*([^\*\*\n]]+)\*\*(?!\w)", "**")
        if len(splitted_runs) == 3:
            splitted_runs[1].bold = True
            runs.append(splitted_runs[1])
            runs.append(splitted_runs[2])
        else:
            splitted_runs = splitRunOnMarker(paragraph, runs[i], r"(?!\w)\_\_([^\_\n]]+)\_\_(?!\w)", "__")
            if len(splitted_runs) == 3:
                splitted_runs[1].bold = True
                runs.append(splitted_runs[1])
                runs.append(splitted_runs[2])
        i+=1
    return runs

def markdownStrikeThroughToStrike(paragraph, initialRun):
    runs = [initialRun]
    i = 0
    while i < len(runs):
        splitted_runs = splitRunOnMarker(paragraph, runs[i], r"(?!\w)\~\~([^\~\n]]+)\~\~(?!\w)", "~~")
        if len(splitted_runs) == 3:
            # THE PYDOCX run.strike does not work
            # Use run.font.strike
            splitted_runs[1].font.strike = True
            runs.append(splitted_runs[1])
            runs.append(splitted_runs[2])
        i+=1
    return runs

def markdownCodeToWordStyle(paragraph, initialRun, style):
    runs = [initialRun]
    i = 0
    while i < len(runs):
        splitted_runs = splitRunOnMarker(paragraph, runs[i], r"(?<!\w)\`([^\`\n]+)\`(?!\w)", "`")
        if len(splitted_runs) == 3:
            # THE PYDOCX run.strike does not work
            # Use run.font.strike
            splitted_runs[1].style = style
            runs.append(splitted_runs[1])
            runs.append(splitted_runs[2])
        i+=1
    return runs
def markdownImgToInsertedImage(paragraph, initialRun):
    runs = [initialRun]
    regex_hyperlink = r"\!\[([^\]|^\n]+)\]\(([^\)|^\n]+\.(?:png|jpg|jpeg|gif))\)"
    regex = re.compile(regex_hyperlink)
    i = 0
    while i < len(runs):
        matched = re.search(regex, runs[i].text)
        if matched is not None:
            start = runs[i].text.index(matched.group(0))
            end = start+len(matched.group(0))
            split_runs = split_run_in_three(paragraph, runs[i], start, end)
            data = downloadImgData(matched.group(2))
            if data is not None:
                split_runs[1].text = ""
                split_runs[1].add_picture(data, width=Cm(17.19))
            else:
                split_runs[1].text = split_runs[1].text.replace(matched.group(0), matched.group(1))
            
        i+=1
    return runs

def markdownLinkToHyperlink(paragraph, initialRun, styles):
    runs = [initialRun]
    regex_hyperlink = r"(?<!\!)\[([^\]|^\n]+)\]\(([^\)|^\n]+)\)"
    regex = re.compile(regex_hyperlink)
    i = 0
    while i < len(runs):
        matched = re.search(regex, runs[i].text)
        if matched is not None:
            start = runs[i].text.index(matched.group(0))
            end = start+len(matched.group(0))
            split_runs = split_run_in_three(paragraph, runs[i], start, end)
            split_runs[1].text = split_runs[1].text.replace(matched.group(0), matched.group(1))
            style = styles["Hyperlink"]
            set_hyperlink(paragraph, split_runs[1], matched.group(2), matched.group(1), style)
        i+=1
    return runs

def splitRunOnMarker(paragraph, run, regexToSearch, markerToRemove):
    regex = re.compile(regexToSearch, re.MULTILINE)
    matched = re.findall(regex, run.text)
    for match in matched:
        start = run.text.index(markerToRemove+match+markerToRemove)
        end = start+len(markerToRemove+match+markerToRemove)
        split_runs = split_run_in_three(paragraph, run, start, end)
        split_runs[1].text = split_runs[1].text.replace(markerToRemove+match+markerToRemove, match)
        return split_runs
    return [run]

def markdownArrayToWordList(document, paragraph, state):
    table_line_regex = re.compile(r"^\|(?:[^\|\n-]*\|)*\s*$", re.MULTILINE)
    matched = re.findall(table_line_regex, paragraph.text)
    if len(matched) == 0:
        return state
    nb_columns = len(matched[0].strip()[1:-1].split("|"))
    array = document.add_table(rows=len(matched), cols=nb_columns)
    for i_row, match in enumerate(matched):
        line = match.strip()
        columns = line[1:-1].split("|") # [1:-1] strip beginning and ending pipe
        if len(columns) != nb_columns:
            raise ValueError("The array with following headers : "+str(matched[0])+" is supposed to have "+str(nb_columns)+ \
                                " columns but the line "+str(line)+" has "+str(len(columns))+" columns")
        for i_column, column in enumerate(columns):
            cell = array.cell(i_row, i_column)
            fill_cell(cell, column)
    move_table_after(array, paragraph)
    delete_paragraph(paragraph)
    return state

def markdownUnorderedListToWordList(paragraph, style, state):
    regex = re.compile(r"^\s*[\*|\-|\+]\s([^\n]+)", re.MULTILINE)
    matched = re.findall(regex, paragraph.text)
    if len(matched) > 0:
        start = paragraph.text.index(matched[0])
        end = paragraph.text.index(matched[-1])+len(matched[-1])
        text_end = paragraph.text[end:]
        paragraph.text = paragraph.text[:start-2].strip() # -2 for list marker + space
        for match in matched:
            new_p = insert_paragraph_after(paragraph)
            new_p.style = "BulletList"
            r = new_p.add_run()
            r.add_text(match)
        if text_end.strip() != "":
            insert_paragraph_after(new_p.strip(), text_end)
        if paragraph.text.strip() == "":
            delete_paragraph(paragraph)
    return state

def mardownCodeBlockToWordStyle(paragraph, styles, state):
    if paragraph.text.lstrip().startswith("```"):
        state = "code_block"
        paragraph.text = paragraph.text.split("```")[0].strip()+"```".join(paragraph.text.split("```")[1:]).strip()
    if state == "code_block":
        paragraph.style = styles["Code"]
    if paragraph.text.rstrip().endswith("```") and state == "code_block":
        state = "normal"
        paragraph.text = "```".join(paragraph.text.split("```")[:-1]).strip()+paragraph.text.split("```")[-1].strip()
    return state

def markdownToWordInParagraph(document, paragraph, styles, state):
    state = markdownArrayToWordList(document, paragraph, state)
    state = markdownUnorderedListToWordList(paragraph, styles["BulletList"], state)
    state = mardownCodeBlockToWordStyle(paragraph, styles, state)
    return state

def markdownToWordInRun(paragraph, initialRun, styles):
    markdownHeaderToWordStyle(paragraph, initialRun, styles["Sous-défaut"])
    new_runs = set(markdownStrongEmphasisToBold(paragraph, initialRun))
    for run in list(new_runs):
        new_runs |= set(markdownEmphasisToItalic(paragraph, run))
    for run in list(new_runs):
        new_runs |= set(markdownStrikeThroughToStrike(paragraph, run))
    for run in list(new_runs):
        new_runs |= set(markdownCodeToWordStyle(paragraph, run, styles["Code Car"]))
    for run in list(new_runs):
        new_runs |= set(markdownImgToInsertedImage(paragraph, run))
    for run in list(new_runs):
        markdownLinkToHyperlink(paragraph, run, styles)

def markdownToWordInDocument(document):
    ps = getParagraphs(document)
    state = "normal"
    for paragraph in ps:
        state = markdownToWordInParagraph(document, paragraph, document.styles, state)
    ps = getParagraphs(document)
    for paragraph in ps:
        for run in paragraph.runs:
            markdownToWordInRun(paragraph, run, document.styles)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        markdownToWordInRun(paragraph, run, document.styles)

def createReport(defects_dict, template, out_name, **kwargs):
    #print("Defect dict: "+str(defects_dict))
    document = Document(template)
    global cell_style
    global normal_style
    global progressbar
    current_step = 0
    cell_style = document.styles["Normal_Cell"]
    normal_style = document.styles["Normal"]
    client_name = kwargs.get("client", "").strip()
    if client_name != "":
        replaceTextInDocument(document, "var_client", client_name)
    date = datetime.now()
    month = date.strftime("%B").lower()
    replaceTextInDocument(document, "var_month", month)
    month_fr = {"january":"janvier", "february":"février", "march":"mars", "april":"avril", "may":"mai", "june":"juin",
                "july":"juillet", "august":"août", "september":"septembre", "october":"octobre", "november":"novembre", "december":"décembre"}
    replaceTextInDocument(document, "var_mois", month_fr.get(month, month))
    replaceTextInDocument(document, "var_year", date.strftime("%Y"))
    replaceTextInDocument(document, "var_annee", date.strftime("%Y"))
    contract_name = kwargs.get("contract", "").strip()
    progressbar = kwargs.get("progressbar", None)
    if progressbar is not None:
        total_len = 0
        levels = ["Critique", "Majeur", "Important", "Mineur"]
        for level in levels:
            total_len += len(defects_dict[level].values())
        nb_steps = total_len # 1 step by defect 
        nb_steps += 1 # step for general stuff
        nb_steps += 1 # step for defect summary
        nb_steps += 1 # step for service table
        nb_steps += 1 # step for markdown conversion
        nb_steps += 1 # step for saving
        progressbar.show(nb_steps)
    else:
        progressbar = DummyProgressBar() # avoid testing for None every time
    if contract_name != "":
        replaceTextInDocument(document, "var_contract", contract_name)
    replaceTextInDocument(document, "var_synthesis", str(kwargs.get("synthesis", "ToDo "+kwargs.get("main_redactor", "synthesis"))))
    replaceTextInDocument(document, "var_nb_d_total", str(len(defects_dict["Critique"].keys())+len(defects_dict["Majeur"].keys())+len(defects_dict["Important"].keys())+len(defects_dict["Mineur"].keys())))
    replaceTextInDocument(document, "var_nb_d_critical", str(len(defects_dict["Critique"].keys())))
    replaceTextInDocument(document, "var_nb_d_major", str(len(defects_dict["Majeur"].keys())))
    replaceTextInDocument(document, "var_nb_d_important", str(len(defects_dict["Important"].keys())))
    replaceTextInDocument(document, "var_nb_d_minor", str(len(defects_dict["Mineur"].keys())))
    progressbar.update() # general stuff step
    print("Populate defect summary ....")
    try:
        populate_defect_summary_table(document, defects_dict)
    except KeyError as e:
        print("Skipping  defect summary: "+str(e))
    progressbar.update() # defect summary step
    print("Write each defect ...")
    write_each_defect(document, defects_dict)
    print("Write services table ...")
    populate_services_table(document, kwargs.get("root", None))
    progressbar.update() # service table step
    print("Converting Markdown ...")
    markdownToWordInDocument(document)
    print("Saving ...")
    # After that, we add the previously copied table
    #paragraph._p.addnext(new_tbl)
    dir_path = os.path.dirname(os.path.realpath(__file__))
    out_path = os.path.join(dir_path, "../../exports/", out_name+".docx")
    print("Generated report at "+str(out_path))
    document.save(out_path)
    progressbar.update() # saving step